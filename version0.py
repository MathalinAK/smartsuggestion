# import os
# from PyPDF2 import PdfReader
# from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
# from langchain_core.documents import Document
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain.vectorstores import Chroma  # Recommended approach
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()

# # 1. PDF to Chunks
# def pdf_to_chunks(pdf_path, chunk_size=1000, chunk_overlap=200):
#     """Directly extract and chunk PDF text"""
#     reader = PdfReader(pdf_path)
#     text = "\n".join([page.extract_text() for page in reader.pages])
    
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=chunk_size,
#         chunk_overlap=chunk_overlap
#     )
#     return splitter.split_text(text)[:5]  # Only first 5 chunks

# # 2. Generate Title from Chunks
# def generate_title(chunks, audience="general public"):
#     """Generate title using Gemini 1.5 Flash"""
#     llm = ChatGoogleGenerativeAI(
#         model="gemini-1.5-flash",
#         google_api_key=os.getenv("GOOGLE_API_KEY"),
#         temperature=0.7
#     )
    
#     combined = "\n\n".join(chunks)
    
#     prompt = f"""
#     Based on these document chunks, create one engaging title for {audience}:
#     {combined}
    
#     The title should:
#     - Be 5-10 words
#     - Capture main themes
#     - Be audience-appropriate
#     - Avoid complex jargon
#     """
    
#     return llm.invoke(prompt).content

# # 3. Main Process
# def process_budget_speech(pdf_path="budget_speech.pdf"):
#     # Step 1: Extract chunks
#     chunks = pdf_to_chunks(pdf_path)
#     print(f"Extracted {len(chunks)} chunks")
    
#     # Step 2: Store in ChromaDB (for potential future use)
#     embeddings = GoogleGenerativeAIEmbeddings(
#         model="models/embedding-001",
#         google_api_key=os.getenv("GOOGLE_API_KEY")
#     )
#     Chroma.from_documents(
#         documents=[Document(page_content=chunk) for chunk in chunks],
#         embedding=embeddings,
#         persist_directory="./budget_db"
#     )
    
#     # Step 3: Generate title
#     title = generate_title(chunks)
#     print("\nGenerated Title:")
#     print(title)

# if __name__ == "__main__":
#     if not os.getenv("GOOGLE_API_KEY"):
#         print("Please create a .env file with GOOGLE_API_KEY")
#     else:
#         process_budget_speech()
import os
import streamlit as st
from PyPDF2 import PdfReader
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from dotenv import load_dotenv
import chromadb
load_dotenv()
google_api_key = os.getenv("GOOGLE_API_KEY")
if not google_api_key:
    st.error("Please create a .env file with your GOOGLE_API_KEY")
    st.stop()

default_states = {
    "keywords": [],
    "title": "",
    "show_keyword_section": False,
    "all_chunks": [],
    "selected_audience": "General Public",
    "analysis_result": "",
    "generated_article": "",
    "refinement_text": "",
    "refined_article": "",
    "current_article": "",
    "generated_post": "",
    "post_type": None,
    "selected_tone": None,
    "custom_tone": ""
}

for key, default_value in default_states.items():
    if key not in st.session_state:
        st.session_state[key] = default_value
st.set_page_config(page_title="AI Post Generator", layout="centered")
st.title("Post Generator")

def pdf_to_limited_chunks(pdf_file, chunk_size=1000, chunk_overlap=200):
    """Extract text from PDF and return only first 5 chunks"""
    try:
        reader = PdfReader(pdf_file)
        text = "\n".join([page.extract_text() for page in reader.pages])
        
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        all_chunks = splitter.split_text(text)
        st.session_state.all_chunks = all_chunks  
        return all_chunks[:5]
    except Exception as e:
        st.error(f"Error processing PDF: {str(e)}")
        return []

def get_llm(temperature=0.5, model="gemini-1.5-flash"):
    """Get LLM instance with specified parameters"""
    return ChatGoogleGenerativeAI(
        model=model,
        google_api_key=google_api_key,
        temperature=temperature
    )

def generate_title(chunks):
    """Generate title using Gemini from limited chunks"""
    if not chunks:
        return ""
        
    llm = get_llm(temperature=0.7)
    combined = "\n\n".join(chunks)
    
    prompt = f"""
    Based on these document chunks, create ONE engaging title:
    {combined}
    
    The title should:
    - Be 5-10 words exactly
    - Capture core themes
    - Avoid complex jargon
    - Include year if mentioned
    
    Return ONLY the title text, nothing else.
    """
    
    response = llm.invoke(prompt).content
    return response.split('\n')[0].strip('"').strip()

def generate_keywords(title, audience):
    """Generate relevant keywords based on title and audience"""
    llm = get_llm(temperature=0.5)
    
    prompt = f"""
    Generate 15 relevant keywords for this title targeting {audience}:
    Title: {title}
    
    The keywords should:
    - Be single words or short phrases
    - Cover main themes
    - Be audience-appropriate
    - Be SEO-friendly
    
    Return ONLY comma-separated keywords.
    """
    
    response = llm.invoke(prompt).content
    keywords = [kw.strip() for kw in response.split(",") if kw.strip()]
    return keywords[:15]

def analyze_keywords(keywords, audience):
    """Analyze how well keywords match the target audience"""
    llm = get_llm(temperature=0.3)
    
    prompt = f"""
    Analyze how well these keywords match the target audience of {audience}:
    Keywords: {", ".join(keywords)}
    
    For each keyword, provide:
    if the keyword matches the target audience, if yes say "yes" else "no" and give a one-line description.
    
    Format your response as a bulleted list with clear separation between keywords.
    """
    
    return llm.invoke(prompt).content

def generate_article(title, keywords, chunks):
    """Generate article based on title and keywords"""
    try:
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=google_api_key
        )
        vector_store = Chroma.from_texts(
            texts=chunks,
            embedding=embeddings,
            collection_name="temp_collection",
            persist_directory=None, 
            client_settings=chromadb.config.Settings(
                anonymized_telemetry=False,
                is_persistent=False
            )
        )
        
        query = f"{title}. Keywords: {', '.join(keywords)}"
        relevant_docs = vector_store.similarity_search(query, k=5)
        relevant_content = "\n\n".join([doc.page_content for doc in relevant_docs])
        
        llm = get_llm(temperature=0.5)
        
        prompt = f"""
        Write one comprehensive, engaging article about: {title}
        ** Easy to Read Make the Article more crisper, more engaging style **
        **make it simple**
        
        CONTENT REQUIREMENTS:
        1. POWERFUL INTRODUCTION
        - Start with a surprising statistic, bold claim, or thought-provoking question
        - Example: "In a groundbreaking development, [shocking fact] about [topic]..."
        
        2. WELL-STRUCTURED BODY
        - Clear subheadings every 2-3 paragraphs
        - Mix of these elements:
        * Big Picture: Industry-wide implications and future outlook
        * Practical Impacts: How this affects businesses/individuals
        * Technical Insights: Simplified explanations of complex aspects
        - Short paragraphs (max 3 sentences)
        - Smooth transitions between sections
        
        3. CONTENT QUALITY
        - Use analogies to explain complex ideas
        - Include 2-3 key statistics/facts
        - Provide real-world examples or case studies
        - Naturally integrate keywords: {', '.join(keywords)}
        
        4. PROFESSIONAL YET ENGAGING TONE
        - Journalistic quality but accessible
        - Avoid excessive jargon
        - Maintain objective perspective
        
        5. STRONG CONCLUSION
        - Summary of key points
        - Future implications
        - Call-to-action or discussion prompt
        
        WORD COUNT: 500-600 words
        
        CONTENT TO REFERENCE:
        {relevant_content}
        """
        
        return llm.invoke(prompt).content
        
    except Exception as e:
        st.error(f"Error generating article: {str(e)}")
        return None

def generate_social_post(article_content, post_type, tone, custom_tone, keywords, audience):
    """Generate social media post based on article content and post type"""
    try:
        selected_tone = tone.split(" ")[0].lower() if tone != "Custom ✏️" else custom_tone.lower()
        temperature = 0.7 if selected_tone == "humorous" else 0.5
        llm = get_llm(temperature=temperature)
        
        post_prompts = {
            "blog": f"""
                Write a **300-400 word blog post** based on this content for {audience}:
                {article_content}
                
                ### Key Guidelines:
                - **Tone:** {selected_tone.upper()} ({custom_tone if tone == "Custom ✏️" else tone})
                - **Hook the Reader** – Start with a bold statement or surprising fact
                - **Engaging Structure** – Use subheadings, bullet points, and short paragraphs(5 lines)
                - **Fresh Insights** – Focus on unique perspectives and real-world impact
                - **Conversational Style** – Keep it {selected_tone} and jargon-free
                - **Credibility** – Back insights with data or examples
                - **SEO Optimization** – Use keywords: {', '.join(keywords)}
                - **Call to Action** – End with a discussion prompt
                
                ### Tone-Specific Enhancements:
                {"- Use emojis and casual language" if selected_tone == "casual" else ""}
                {"- Maintain professional terminology" if selected_tone == "formal" else ""}
                {"- Include tasteful humor and wit" if selected_tone == "humorous" else ""}
                {"- Follow custom tone description exactly" if tone == "Custom " else ""}
                
                Return ONLY the formatted blog post.
                """,
                
            "linkedin": f"""
                Write a compelling **300-word LinkedIn post** based on this content that grabs attention and makes people *stop scrolling*. 💥 Ensure it's **engaging**, **thought-provoking**, and **encourages interaction**.
                {article_content}
                
                **Requirements:**
                - Tone: {selected_tone.upper()} ({custom_tone if tone == "Custom ✏️" else tone})
           
                -**Hook:** Start with a bold statement, surprising fact, or a relatable question (1 line).\n
                -**Make It Skimmable:** Use short sentences, line breaks, and bold key points for easy reading.\n
                -**Explain Simply:** Describe the concept in a crisp, easy-to-understand way (1 line).\n
                -**Add a Quick Analogy or Example:** Make it relatable with a simple comparison (1 line).\n
                -**Call to Action:** End with a thought-provoking question to spark discussion (1 line).\n
                -**Use Hashtags:** Add relevant hashtags at the end.\n\n
                -**Tone:**\n
                -**content has to be like human written and more crisp short etc**
                - Engaging, simple, and beginner-friendly.\n
                - Short, punchy, and easy to skim.\n
                - Informative with a touch of wit—just enough to make it interesting!
                        
                **Tone Guidelines:**
                {"- Casual, friendly, with emojis" if selected_tone == "casual" else ""}
                {"- Professional but engaging" if selected_tone == "formal" else ""}
                {"- Witty and humorous" if selected_tone == "humorous" else ""}
                {"- Custom: " + custom_tone if tone == "Custom ✏️" else ""}
                
                Return ONLY the LinkedIn post content.
                """,
                
            "twitter": f"""
                Write an engaging tweet thread based on this content:
                {article_content}
                
                **Requirements:**
                - Tone: {selected_tone.upper()} ({custom_tone if tone == "Custom ✏️" else tone})
                - Keep it **under 240 chars**, natural tone (no all caps).  
                - **Front-load key message** in the first 3-4 words.  
                - Spark **emotion, passion, or excitement**.  
                - Add a **clear CTA** (reply, click, share).  
                - Use **2-5 hashtags & 1 emoji** for reach.  
                - Tag someone or add a **link** if relevant.  
                - Ensure it's **engaging & optimized for interaction**.
                
                **Tone Guidelines:**
                {"- Casual, conversational" if selected_tone == "casual" else ""}
                {"- Professional but concise" if selected_tone == "formal" else ""}
                {"- Humorous and playful" if selected_tone == "humorous" else ""}
                {"- Custom: " + custom_tone if tone == "Custom ✏️" else ""}
                """,
                
            "email": f"""
                Write a professional email based on this content:
                {article_content}
                
                **Requirements:**
                - Tone: {selected_tone.upper()} ({custom_tone if tone == "Custom ✏️" else tone})
                -**Make It More Personalized eg("  Dear [Name],")
                Write an engaging, **scannable email** with:
                - **Compelling hook** (question or bold statement)
                - **Short paragraphs & bullet points** for readability
                - **Inverted pyramid structure** leading to a strong **CTA**
                Use **keywords naturally**:  avoiding repetition.  
                Ensure a **smooth, conversational flow**.  
                Write a **clear, action-driven CTA** that encourages interaction.  
                Optimize for **mobile readability** & **avoid spam triggers**.  
                Keep it **engaging, relevant & thought-provoking**.
                
                **Tone Guidelines:**
                {"- Friendly and approachable" if selected_tone == "casual" else ""}
                {"- Formal and professional" if selected_tone == "formal" else ""}
                {"- Lighthearted with humor" if selected_tone == "humorous" else ""}
                {"- Custom: " + custom_tone if tone == "Custom ✏️" else ""}
                
                Format:
                Subject: [subject line]
                
                [email body]
                
                Return ONLY the email content.
                """
        }
        
        prompt = post_prompts.get(post_type, "")
        return llm.invoke(prompt).content if prompt else None
        
    except Exception as e:
        st.error(f"Error generating post: {str(e)}")
        return None

def refine_article(current_article, refinement_instruction, keywords):
    """Refine article based on user instructions"""
    try:
        llm = get_llm(temperature=0.4)
        
        refine_prompt = f"""
        Please refine the following article based on these specific instructions:
        
        REFINEMENT REQUEST:
        {refinement_instruction}
        
        CURRENT ARTICLE CONTENT:
        {current_article}
        
        GUIDELINES FOR REFINEMENT:
        1. Make only the requested changes - don't modify other parts
        2. Keep the same overall structure and tone
        3. Maintain all key facts and information
        4. Preserve the keyword integration: {', '.join(keywords)}
        5. Highlight changes by bolding new or modified text
        
        OUTPUT REQUIREMENTS:
        - Return the complete revised article
        - Mark changes in bold
        - Keep the same approximate length
        - Maintain all original section headings
        """
        
        return llm.invoke(refine_prompt).content
    except Exception as e:
        st.error(f"Error refining article: {str(e)}")
        return None

def reset_state_after(state_to_keep):
    """Reset state variables after certain operations"""
    states_to_reset = {
        "title": ["analysis_result", "generated_article", "refined_article", 
                 "current_article", "generated_post"],
        "keywords": ["analysis_result", "generated_article", "refined_article", 
                    "current_article", "generated_post"],
    }
    
    for state_to_reset in states_to_reset.get(state_to_keep, []):
        st.session_state[state_to_reset] = ""
    
    if state_to_keep == "keywords":
        st.session_state.post_type = None
uploaded_file = st.file_uploader("Upload the document", type=["pdf"])

if uploaded_file is not None:
 
    if st.button("Generate Title"):
        with st.spinner("Analyzing document..."):
            try:
                selected_chunks = pdf_to_limited_chunks(uploaded_file)
                if selected_chunks:
                    st.session_state.title = generate_title(selected_chunks)
                    st.session_state.show_keyword_section = True
                    reset_state_after("title")
                    st.rerun()
            except Exception as e:
                st.error(f"Error generating title: {str(e)}")
    
  
    if st.session_state.title:
        st.markdown(f"## {st.session_state.title}")
        
        if st.session_state.show_keyword_section:
            st.divider()
            st.session_state.selected_audience = st.selectbox(
                "Select Target Audience",
                ["General Public", "Business Leaders", "Policy Makers", 
                 "Investors", "Media", "Sales Teams", "Marketing Professionals"],
                index=0
            )
            
            # Generate keywords
            if st.button("Generate Keywords"):
                with st.spinner("Generating keywords..."):
                    st.session_state.keywords = generate_keywords(
                        st.session_state.title, 
                        st.session_state.selected_audience
                    )
                    reset_state_after("keywords")
                    st.rerun()
            
            # Display and manage keywords
            if st.session_state.keywords:
                st.subheader("Keywords")
                cols = st.columns(4)
                keywords_to_remove = []
                
                for i, kw in enumerate(st.session_state.keywords):
                    with cols[i % 4]: 
                        if st.button(f"ˣ {kw}", key=f"del_{kw}"):
                            keywords_to_remove.append(kw) 
                
                if keywords_to_remove:
                    st.session_state.keywords = [kw for kw in st.session_state.keywords 
                                               if kw not in keywords_to_remove]
                    reset_state_after("keywords")
                    st.rerun()
                custom_key = f"custom_keyword_{len(st.session_state.keywords)}"
                custom_keyword = st.text_input(
                    "Add custom keyword", 
                    key=custom_key, 
                    placeholder="Enter a keyword"
                )
                
                if st.button("Add Keyword"):
                    if custom_keyword.strip() and custom_keyword.strip() not in st.session_state.keywords:
                        st.session_state.keywords.append(custom_keyword.strip())
                        reset_state_after("keywords")
                        st.rerun()
                    elif custom_keyword.strip() in st.session_state.keywords:
                        st.warning("Keyword already exists in the list")
                
                st.divider()
                
                # Analyze keywords
                if st.button("Analyze Keyword Relevance"):
                    with st.spinner("Analyzing keyword relevance..."):
                        st.session_state.analysis_result = analyze_keywords(
                            st.session_state.keywords,
                            st.session_state.selected_audience
                        )
                        st.rerun()
                
                if st.session_state.analysis_result:
                    st.subheader("Keyword Relevance Analysis")
                    st.markdown(st.session_state.analysis_result)
                    st.divider()
                
                # Generate article
                if st.button("Generate Article"):
                    with st.spinner("Creating article from document..."):
                        article = generate_article(
                            st.session_state.title,
                            st.session_state.keywords,
                            st.session_state.all_chunks
                        )
                        if article:
                            st.session_state.generated_article = article
                            st.session_state.current_article = article
                            st.session_state.refined_article = ""
                            st.rerun()

                if st.session_state.generated_article:
                    st.subheader("Generated Article")
                    st.markdown(st.session_state.current_article)
                    st.divider()
                    st.subheader("Refine Article")
                    with st.form("refinement_form"):
                        user_input = st.text_input(
                            "Enter refinement instructions",
                            value=st.session_state.refinement_text,
                            placeholder="What would you like to change in the article?",
                            key=f"refinement_input_{hash(st.session_state.current_article)}"
                        )
                        submitted = st.form_submit_button("Refine Article")
                    
                    if submitted and user_input.strip():
                        with st.spinner("Refining article..."):
                            refined = refine_article(
                                st.session_state.current_article,
                                user_input,
                                st.session_state.keywords
                            )
                            if refined:
                                st.session_state.refined_article = refined
                                st.session_state.current_article = refined
                                st.session_state.refinement_text = ""
                                st.rerun()
                    elif submitted:
                        st.warning("Please enter refinement instructions")
                    else:
                        st.session_state.refinement_text = user_input
                    
                    # Article selection
                    if st.session_state.refined_article:
                        st.divider()
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.button("Use Original Article"):
                                st.session_state.current_article = st.session_state.generated_article
                                st.rerun()
                        with col2:
                            if st.button("Use Refined Article"):
                                st.session_state.current_article = st.session_state.refined_article
                                st.rerun()
                #post type
                    st.divider()
                    st.subheader("Convert to Social Post")
                    st.subheader("Select Post Type")
                    col1, col2, col3, col4 = st.columns(4)
                    
                    post_types = {
                        "blog": "Blog Post",
                        "linkedin": "LinkedIn Post",
                        "twitter": "Twitter",
                        "email": "Email Post"
                    }
                    
                    cols = [col1, col2, col3, col4]
                    for i, (key, label) in enumerate(post_types.items()):
                        with cols[i]:
                            if st.button(f" {label}", key=f"{key}_post_btn", use_container_width=True):
                                st.session_state.post_type = key
                                st.rerun()
                    if st.session_state.post_type:
                        st.subheader("Select Tone")
                        tone = st.radio(
                            "Choose writing style:",
                            ["Casual ", "Formal ", "Humorous ", "Custom "],
                            horizontal=True,
                            key="tone_selection",
                            label_visibility="collapsed"
                        )
                        
                        st.session_state.selected_tone = tone
                    
                        if tone == "Custom ✏️":
                            custom_tone = st.text_input(
                                "Describe the tone you want (e.g., 'Motivational', 'Technical', 'Friendly expert')", 
                                key="custom_tone_input"
                            )
                            st.session_state.custom_tone = custom_tone
                        
                        # Generate social post
                        if st.button(f"Generate {st.session_state.post_type.replace('_', ' ').title()} Post", 
                                   key="generate_post_btn"):
                            with st.spinner(f"Crafting your {st.session_state.post_type} post..."):
                                post_content = generate_social_post(
                                    st.session_state.current_article,
                                    st.session_state.post_type,
                                    st.session_state.selected_tone,
                                    st.session_state.custom_tone,
                                    st.session_state.keywords,
                                    st.session_state.selected_audience
                                )
                                
                                if post_content:
                                    st.session_state.generated_post = post_content
                                    st.rerun()
                        if st.session_state.generated_post:
                            st.subheader(f"Your {st.session_state.post_type.title()} Post")
                            st.markdown(st.session_state.generated_post)



import streamlit as st
import google.generativeai as genai
import fitz  
from docx import Document
from dotenv import load_dotenv
import os
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=GOOGLE_API_KEY)

def extract_text_from_pdf(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = "\n".join([page.get_text("text") for page in doc])
    return text

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def get_gemini_title(text):
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = (
        f"Based on the following content, generate a **catchy, engaging, and well-structured** title that "
        f"grabs attention and clearly represents the main idea:\n\n{text[:2000]}\n\n"
        "Ensure the title is **strictly one**, clear, engaging, and accurately represents the main idea."
    )
    response = model.generate_content(prompt)
    return response.text.strip()
def get_gemini_keywords(title, text):
    model = genai.GenerativeModel("gemini-1.5-flash")
    
    prompt = (
        f"Extract **15 to 20 highly relevant keywords** from the following content:\n\n"
        f"Title: {title}\n\nContent:\n{text[:2000]}\n\n"
        "- Focus on **key concepts, technical terms, industry-specific words, and important ideas**.\n"
        "- **Ignore common words, numbers, dates, locations, and generic terms** (e.g., '12th', 'Nepal', 'Kailali').\n"
        "- Provide the keywords as a **comma-separated list** without numbering.\n"
        "- Ensure at least **15 keywords** are extracted. If fewer than 15 are found, use synonyms or related terms."
    )
    
    response = model.generate_content(prompt)
    keywords = list(set(response.text.strip().split(", ")))
    if len(keywords) < 15:
        prompt += "\nReturn exactly **15-20 keywords** by expanding on key ideas, synonyms, and technical terms."
        response = model.generate_content(prompt)
        keywords = list(set(response.text.strip().split(", ")))

    return keywords[:20]  

# Generate three well-structured articles
def generate_articles():
    keywords_str = ", ".join(st.session_state.keywords)
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"""Write **three engaging and well-structured articles** on **{st.session_state.title}**.  
    **1. Hook the Reader**  
    Start with a **bold claim, surprising fact, or compelling question** to grab attention instantly.  
    **2. Make It Relatable**  
    Use **analogies, storytelling, or real-world examples** to simplify complex ideas.  
    **3. Ensure Readability**  
    - Use **clear subheadings** to organize content.  
    - Keep paragraphs **short and to the point** for easy skimming.  
    - Maintain **smooth transitions** between sections.  
    **4. Simplify Complexity**  
    - Break down **abstract ideas** using analogies.  
    - Use **research-backed insights** to add credibility.  
    **5. Balance Depth & Clarity**  
    - Keep the tone **conversational yet professional** (avoid jargon overload).  
    - Offer **actionable takeaways** that provide value.  
    - Naturally integrate **{keywords_str}** into the content.  
    **6. Drive Action**  
    - Highlight the **real-world impact** of the topic.  
    - Discuss **future implications** and industry trends.  
    - End with a **strong call to action** to spark engagement.  
     **7. Keep It Concise**  
    Each article should be **under 500 words** for maximum engagement.  
    **Formatting Tip:** Separate each article with:  
    `---ARTICLE END---`
    """
    response = model.generate_content(prompt)
    articles = response.text.strip().split("---ARTICLE END---")
    st.session_state.articles = [article.strip() for article in articles if article.strip()]
    st.rerun()

# Generate content based on the selected article 
def generate_content(content_type, selected_article):
    keywords_str = ", ".join(st.session_state.keywords)
    model = genai.GenerativeModel("gemini-1.5-flash")
    if content_type == "Blog":
        prompt = f"""
        Write a **300-400 word blog post** for your target audience on **{st.session_state.title}**.
        ### Key Guidelines:
        - **Hook the Reader** – Start with a bold statement or surprising fact to grab attention.
        - **Engaging & Clear Structure** – Use **subheadings, bullet points, and short paragraphs** for readability.
        - **Fresh Insights** – Focus on unique perspectives, emerging trends, and real-world impact.
        - **Conversational Tone** – Keep it **friendly, direct, and jargon-free**, addressing the reader as “you.”
        - **Credibility & Examples** – Back insights with **data, expert opinions, or real-life examples** to add authority.
        - **SEO Optimization** – Naturally integrate relevant **keywords** and **internal links** for better visibility.
        - **Call to Action** – End with a strong **question or discussion prompt** to encourage engagement.
        ### Enhancements for Impact:
        1. **Stronger Headings** - Keep them concise, clear, and engaging.
        2. **Smooth Transitions** – Ensure a natural flow between sections.
        3. **Compelling CTA** – Encourage readers to **share opinions and join the discussion.**
        4. **Relatable Examples** – Use real-life scenarios to make complex topics accessible.
        Make it insightful, engaging, and easy to read!
        """
    elif content_type == "LinkedIn":
       prompt = (
        f"Write a compelling **300-word** LinkedIn post on **{st.session_state.title}** that grabs attention and makes people *stop scrolling*. 💥 Ensure it's **engaging**, **thought-provoking**, and **encourages interaction**."
        f" **Hook:** Start with a bold statement, surprising fact, or a relatable question (1 line).\n"
        f" **Make It Skimmable:** Use short sentences, line breaks, and bold key points for easy reading.\n"
        f" **Explain Simply:** Describe the concept in a crisp, easy-to-understand way (1 line).\n"
        f" **Add a Quick Analogy or Example:** Make it relatable with a simple comparison (1 line).\n"
        f" **Call to Action:** End with a thought-provoking question to spark discussion (1 line).\n"
        f" **Use Hashtags:** Add relevant hashtags at the end.\n\n"
        f" **Tone:**\n"
        f"content has to be like human written and more crisp short etc"
        f"- Engaging, simple, and beginner-friendly.\n"
        f"- Short, punchy, and easy to skim.\n"
        f"- Informative with a touch of wit—just enough to make it interesting!"
    )
    elif content_type == "Twitter":
        prompt = f"""
        Write a short, engaging tweet on **{st.session_state.title}** (max 280 chars).  
        - Keep it **under 240 chars**, natural tone (no all caps).  
        - **Front-load key message** in the first 3-4 words.  
        - Spark **emotion, passion, or excitement**.  
        - Add a **clear CTA** (reply, click, share).  
        - Use **2-5 hashtags & 1 emoji** for reach.  
        - Tag someone or add a **link** if relevant.  
        - Ensure it's **engaging & optimized for interaction**.
        """
    elif content_type == "Email":
        prompt =f"""
        Generate a **catchy subject line** for **{st.session_state.title}**.
        Write an engaging, **scannable email** with:
        - **Compelling hook** (question or bold statement)
        - **Short paragraphs & bullet points** for readability
        - **Inverted pyramid structure** leading to a strong **CTA**
        Use **keywords naturally**: {keywords_str}, avoiding repetition.  
        Ensure a **smooth, conversational flow**.  
        Write a **clear, action-driven CTA** that encourages interaction.  
        Optimize for **mobile readability** & **avoid spam triggers**.  
        Keep it **engaging, relevant & thought-provoking**.
        """ 
    response = model.generate_content(prompt)
    st.session_state.generated_content = response.text.strip()
    st.rerun()
st.title("AI-Powered Article Generator")
uploaded_file = st.file_uploader("Upload a PDF or DOCX file", type=["pdf", "docx"])
if "title" not in st.session_state:
    st.session_state.title = None
if "keywords" not in st.session_state:
    st.session_state.keywords = []
if "articles" not in st.session_state:
    st.session_state.articles = []
if "generated_content" not in st.session_state:
    st.session_state.generated_content = ""
if "prev_uploaded_file" not in st.session_state:
    st.session_state.prev_uploaded_file = None

if uploaded_file is not None:
    if uploaded_file.name != st.session_state.prev_uploaded_file:
        st.session_state.prev_uploaded_file = uploaded_file.name  

        with st.spinner("Extracting text..."):
            if uploaded_file.name.endswith(".pdf"):
                extracted_text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.name.endswith(".docx"):
                extracted_text = extract_text_from_docx(uploaded_file)
            else:
                st.error("Unsupported file type. Please upload a .pdf or .docx file.")
                extracted_text = ""

        if extracted_text:
            with st.spinner("Generating title and keywords..."):
                title = get_gemini_title(extracted_text)
                keywords = get_gemini_keywords(title, extracted_text) 

            st.session_state.title = title
            st.session_state.keywords = keywords


    if st.session_state.title:
        st.subheader("Suggested Title:")
        st.write(f"**{st.session_state.title}**")

    st.subheader("Suggested Keywords:")
    
    cols = st.columns(4)
    keywords_to_remove = []

    for i, kw in enumerate(st.session_state.keywords):
        with cols[i % 4]: 
            if st.button(f"ˣ {kw}", key=f"del_{kw}"):
                keywords_to_remove.append(kw) 
                
    if keywords_to_remove:
        st.session_state.keywords = [kw for kw in st.session_state.keywords if kw not in keywords_to_remove]
        st.rerun()

    custom_key = f"custom_keyword_{len(st.session_state.keywords)}"
    custom_keyword = st.text_input("Add custom keyword", key=custom_key, placeholder="Enter a keyword")

    if st.button("Add Keyword"):
        if custom_keyword.strip() and custom_keyword.strip() not in st.session_state.keywords:
            st.session_state.keywords.append(custom_keyword.strip())  
            st.rerun()

    st.subheader("Generate Articles")
    if st.button("Generate Articles"):
        generate_articles()
if "selected_index" not in st.session_state:
    st.session_state.selected_index = None
if "selected_article" not in st.session_state:
    st.session_state.selected_article = None

if st.session_state.articles:
    st.subheader("Generated Articles:")

    for i, article in enumerate(st.session_state.articles):
        title = article.splitlines()[0] 
        is_selected = st.session_state.selected_index == i
        article_style = "background-color: #F1C1C0; padding: 7px; border-radius: 5px;" if is_selected else ""
        cols = st.columns([0.05, 0.95])
        with cols[0]:
            if st.button("🔴" if is_selected else "⚪", key=f"select_{i}"):
                st.session_state.selected_index = i 
                st.session_state.selected_article = article
                st.rerun()  
        with cols[1]:
            st.markdown(f"<div style='{article_style}'><b>{title}</b></div>", unsafe_allow_html=True)
        st.write(article)
        st.divider()  
    if st.session_state.selected_article:
        st.subheader("Generate Content")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("Blog Content"):
                generate_content("Blog", st.session_state.selected_article)
        with col2:
            if st.button("LinkedIn Content"):
                generate_content("LinkedIn", st.session_state.selected_article)
        with col3:
            if st.button("Twitter Content"):
                generate_content("Twitter", st.session_state.selected_article)
        with col4:
            if st.button("Email Content"):
                generate_content("Email", st.session_state.selected_article)
        if st.session_state.get("generated_content"):
            st.subheader("Generated Content:")
            st.write(st.session_state.generated_content)
